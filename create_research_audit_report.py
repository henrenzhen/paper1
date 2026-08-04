from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "deliverables"
OUT_PATH = OUT_DIR / "ZDS论文拒稿复盘与SCI_Q2+重构路线_证据审计版.docx"


NAVY = "16324F"
BLUE = "2E5D8A"
TEAL = "177E89"
RED = "B64040"
AMBER = "B77816"
GREEN = "2E7D5B"
LIGHT_BLUE = "EAF1F7"
LIGHT_TEAL = "E8F4F3"
LIGHT_RED = "F9EAEA"
LIGHT_AMBER = "FFF4DB"
LIGHT_GREEN = "EAF5EF"
LIGHT_GREY = "F2F4F6"
MID_GREY = "D9DEE3"
DARK = "22303C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn("w:" + key), str(edge_data[key]))


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_header(row) -> None:
    set_repeat_table_header(row)
    prevent_row_split(row)


def set_table_widths(table, widths_cm: Sequence[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def style_table_text(table, size=8.3) -> None:
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": MID_GREY},
                bottom={"val": "single", "sz": 4, "color": MID_GREY},
                left={"val": "single", "sz": 4, "color": MID_GREY},
                right={"val": "single", "sz": 4, "color": MID_GREY},
            )
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(size)
                    run.font.color.rgb = RGBColor.from_string(DARK)
        if row_index == 0:
            set_repeat_header(row)
            for cell in row.cells:
                set_cell_shading(cell, NAVY)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths=None, size=8.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for r_idx, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for c_idx, value in enumerate(values):
            cells[c_idx].text = str(value)
        if r_idx % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, "F8FAFB")
    if widths:
        set_table_widths(table, widths)
    style_table_text(table, size=size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def add_hyperlink(paragraph, text: str, url: str, color=BLUE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_field(run, field_code: str) -> None:
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_callout(doc: Document, title: str, body: str, fill=LIGHT_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(
        cell,
        left={"val": "single", "sz": 18, "color": accent},
        top={"val": "nil", "sz": 0, "color": fill},
        bottom={"val": "nil", "sz": 0, "color": fill},
        right={"val": "nil", "sz": 0, "color": fill},
    )
    set_cell_margins(cell, top=130, bottom=130, start=170, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.2, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc: Document, text: str, level=0, bold_prefix: str | None = None):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=9.5, bold=True, color=DARK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=9.5, color=DARK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=DARK)
    return p


def add_number(doc: Document, text: str, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=DARK)
    return p


def add_body(doc: Document, text: str, bold_prefix: str | None = None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=9.6, color=DARK, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=9.6, color=DARK, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=9.6, color=DARK, italic=italic)
    return p


def add_code_line(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(5)
    p.paragraph_format.right_indent = Mm(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run_font(r, name="Consolas", size=8.2, color=NAVY)
    return p


def add_section_title(doc: Document, number: str, title: str, subtitle: str | None = None):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.keep_with_next = True
    r1 = p.add_run(number + "  ")
    set_run_font(r1, size=19, color=TEAL, bold=True)
    r2 = p.add_run(title)
    set_run_font(r2, size=19, color=NAVY, bold=True)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        r = p2.add_run(subtitle)
        set_run_font(r, size=9.2, color="5E6B75", italic=True)


def add_subtitle(doc: Document, title: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, size=12.5, color=BLUE, bold=True)
    return p


def add_review(doc: Document, verdict: str, checks: Iterable[str], risk: str):
    add_callout(doc, "阶段 Review · " + verdict, risk, fill=LIGHT_GREEN, accent=GREEN)
    for item in checks:
        add_bullet(doc, "✓ " + item)


def style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in (("Heading 1", 19, NAVY), ("Heading 2", 12.5, BLUE), ("Heading 3", 10.5, TEAL)):
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.space_before = Pt(9)
        st.paragraph_format.space_after = Pt(4)

    for name in ("List Bullet", "List Bullet 2", "List Number", "List Number 2"):
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(9.5)
        st.paragraph_format.space_after = Pt(2)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(16)
        section.left_margin = Mm(18)
        section.right_margin = Mm(18)
        section.header_distance = Mm(7)
        section.footer_distance = Mm(7)

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("ZDS 论文证据审计  ·  2026-07-29")
        set_run_font(r, size=7.8, color="7C8994")

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run("内部研究决策材料   |   ")
        set_run_font(r1, size=7.5, color="7C8994")
        r2 = p.add_run()
        set_run_font(r2, size=7.5, color="7C8994")
        add_field(r2, "PAGE")


def add_manual_contents(doc: Document) -> None:
    add_section_title(doc, "导读", "报告结构与证据等级")
    add_callout(
        doc,
        "一句话结论",
        "这次两天拒稿首先是 Computers & Security 的当前 AI/ML 范围禁令所致；但现稿即使转投其他期刊，也必须先修复实验溯源和数据划分，再把贡献升级为“可靠性 + OOD 评测协议”，而不是继续包装固定权重融合。",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )
    add_table(
        doc,
        ["章节", "回答的问题", "最终产出"],
        [
            ("0", "核心判断是什么？", "去留决策与 P0 门槛"),
            ("1", "审计覆盖了哪些材料？", "证据等级与限制"),
            ("2 / 任务1", "为什么秒拒、创新为何不足？", "分层根因与代码证据"),
            ("3 / 任务2", "下一篇应做什么创新？", "主路线 + 两条备选路线"),
            ("4 / 任务3", "是否与既有论文重复？", "20 篇近邻与差异边界"),
            ("5 / 任务4", "实验数据从哪里来、能否用？", "A–D 分级与最小数据组合"),
            ("6", "怎样落地为 Q2+ 稿件？", "实验协议、时间表、投稿路径"),
            ("附录", "如何复核全部结论？", "代码证据索引、文献与链接"),
        ],
        widths=[2.2, 7.0, 7.4],
        size=8.8,
    )
    add_body(doc, "证据等级：A=文件/源码/结构化产物或官方页面直接支持；B=多项证据一致但仍需原始日志/拒稿信确认；C=方法学判断或未来方案，不能写成既成事实。")


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("证据审计版")
    set_run_font(r, size=11, color=WHITE, bold=True)
    set_cell = doc.add_table(rows=1, cols=1)
    set_cell.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = set_cell.cell(0, 0)
    set_cell_shading(c, TEAL)
    set_cell_margins(c, top=80, bottom=80, start=220, end=220)
    c.paragraphs[0]._p.addnext(p._p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("ZDS 论文拒稿复盘与\nSCI Q2+ 重构路线")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于论文 PDF、源码、实验产物、期刊政策、近邻文献与公开数据集的联合审计")
    set_run_font(r, size=12, color=BLUE)

    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [
        ("审计对象", "Enhancing ATT&CK Parent-Technique Next-Step Prediction with Controlled LLM-Derived Semantic Fusion"),
        ("论文", r"E:\desktop\ZDS PAPER.pdf"),
        ("源码与数据", r"E:\desktop\project_only\project"),
        ("审计日期", "2026-07-29（Asia/Shanghai）"),
    ]
    for i, (k, v) in enumerate(values):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
    set_table_widths(table, [3.2, 12.8])
    style_table_text(table, 8.8)

    doc.add_paragraph()
    add_callout(
        doc,
        "使用边界",
        "本报告用于科研决策和重现实验，不是对作者动机的判断。凡仓库缺少原始日志或构造脚本之处，只下“不可验证/无实验溯源”的结论，不推断主观行为。期刊拒稿信未提供，因此“两天拒稿”原因是高置信推断，而非编辑部事实陈述。",
        fill=LIGHT_GREY,
        accent="697782",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("Prepared for research decision-making  ·  Confidential draft")
    set_run_font(r, size=8.5, color="7C8994", italic=True)
    doc.add_page_break()


def section_executive(doc: Document) -> None:
    add_section_title(doc, "0", "执行摘要", "先给决策，再给证据")
    add_table(
        doc,
        ["判断", "结论", "置信度", "建议"],
        [
            ("两天拒稿主因", "当前稿以 GRU/LLM/BGE/MLP 为主体，直接触发 C&S 自 2024 年初实施的 AI/ML 稿件暂停考虑政策 [S1]。", "很高", "不要以当前方法重投 C&S"),
            ("现稿创新性", "任务和组件均已有强近邻；固定 α late fusion 是标准组合，尚不足以单独支撑 Q2+。", "高", "把主张从“融合提高准确率”升级为“可靠选择性预测”"),
            ("现稿证据链", "表3来自源码中明确标注为示例的手填数组；表2无对应产物；论文与代码存在多处方法不一致。", "很高", "撤下现有表2/3，全部自动重跑"),
            ("数据可信度", "样本量可核实；sequence_id 无重叠；campaign 隔离无字段/manifest 可证，SIM 根组大量跨 split。", "高", "重建 actor/campaign/report/time 分组切分"),
            ("最优新路线", "分歧感知 conformal selective fusion + actor/time/ATT&CK-version OOD benchmark。", "中高", "作为一篇完整新论文，而非修辞性补丁"),
            ("可用公开数据", "Attack Flow 与 AEL 是核心序列资源；ATT&CK STIX 用于版本/标签本体；TIE/TRAM/AnnoCTR 只作辅助。", "高", "采用来源隔离和泄漏清洗"),
        ],
        widths=[2.7, 8.2, 1.8, 4.7],
        size=8.2,
    )
    add_callout(
        doc,
        "Go / No-Go",
        "NO-GO：当前 PDF 不应直接转投。GO 的最低门槛是：①重建可审计 split；②从 checkpoint/逐样本预测自动生成全部表；③修复 preprocessing 与 No-CoT 对齐；④完成 cluster-aware 统计；⑤主创新改为不确定性、拒识和真实 OOD，而不是固定 α。",
        fill=LIGHT_RED,
        accent=RED,
    )
    add_subtitle(doc, "目标论文的一句话定位")
    add_body(doc, "建议英文定位：Reliable ATT&CK Next-Technique Forecasting under Threat-Actor and Knowledge-Base Shift via Disagreement-Aware Conformal Selective Fusion。")
    add_body(doc, "这一定位保留你现有的 GRU 序列分支、LLM 语义分支和 ATT&CK parent-technique 任务，但贡献中心从“准确率小幅提升”转向“在冲突、长尾和分布漂移下知道何时预测、给出多大候选集合、何时拒绝”。")


def section_scope(doc: Document) -> None:
    doc.add_page_break()
    add_section_title(doc, "1", "材料、方法与证据边界")
    add_subtitle(doc, "1.1 已审计材料")
    add_bullet(doc, "论文 PDF：13 页；全文抽取并将每页渲染为图片逐页视觉检查。")
    add_bullet(doc, "项目源码：rl、llm、fusion、adaptive fusion、Ablation experiment、data、data_v2 等目录。")
    add_bullet(doc, "核心数据：train/val/test 共 14,128 个前缀样本，184 个 parent-technique 标签。")
    add_bullet(doc, "结构化产物：逐样本预测、校准、成本、覆盖率、多种子汇总、外部 CTID 与 micro-state 结果。")
    add_bullet(doc, "外部证据：期刊官方 Guide for Authors、出版商/DOI 页面、官方 GitHub/机构数据页。")

    add_subtitle(doc, "1.2 审计方法")
    add_number(doc, "论文主张 → 源码实现 → 数据/产物三向对照。")
    add_number(doc, "对核心 CSV 重算规模、重叠、标签频次和现存预测指标。")
    add_number(doc, "对 Top-1 四位小数执行“整数命中数 / 2,107”可实现性检查。")
    add_number(doc, "以官方页面为准核验期刊范围；以出版商/DOI/作者全文核验近邻论文。")
    add_number(doc, "数据集按整文件解析、目录/HEAD、小样本、仅元数据四种验证深度分级。")

    add_subtitle(doc, "1.3 限制")
    add_bullet(doc, "未收到编辑拒稿信，因此不能确认编辑给出的正式理由；C&S 范围冲突是基于当前官方政策和拒稿时长的高置信解释。")
    add_bullet(doc, "当前环境缺少可直接运行这些 checkpoint 的 PyTorch 环境；本报告没有重做五种子前向推理，只审计源码、checkpoint 存在性和现有产物。")
    add_bullet(doc, "大型 OpTC/LANL/Attack Flow/AEL 未全量下载；报告明确区分“完整解析”“HEAD/目录验证”和“官方规模声明”。")
    add_bullet(doc, "搜索未发现完全重复论文不等于绝对不存在；结论限定为“截至 2026-07-29 的官方/出版商可检索证据”。")

    add_review(
        doc,
        "通过",
        [
            "PDF 全 13 页已检查，无缺页；关键表格、公式、图和参考文献均纳入审计。",
            "源码结论均保留文件路径/行号或结构化产物作为复核入口。",
            "外部检索只采用官方、出版商、DOI、作者全文或机构数据页。",
        ],
        "材料覆盖足以回答五项任务；不能替代重新运行完整训练和取得拒稿信。",
    )


def section_task1(doc: Document) -> None:
    doc.add_page_break()
    add_section_title(doc, "2", "任务1｜为什么创新不足，为什么会被秒拒")
    add_callout(
        doc,
        "核心结论",
        "“秒拒”主要不是审稿人用两天否定了你的模型，而是投稿范围在编辑初筛阶段就冲突；与此同时，当前版本确实存在会让其他 Q2+ 期刊拒绝的创新、证据链和复现问题。两者必须分开。",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )

    add_subtitle(doc, "2.1 第一层原因：期刊范围硬冲突")
    add_body(doc, "Computers & Security 的当前 Guide for Authors 明确写明：自 2024 年初起，该刊暂停考虑以 AI 或 ML 为重要组成部分的投稿；将 AI/ML 应用于系统安全和隐私主题的稿件不予考虑，并建议投向以 AI/ML 为主要方向的期刊 [S1]。你的标题、摘要、方法和贡献均以 GRU、LLM、BGE 与 logit fusion 为中心，属于政策直接命中的类型。")
    add_body(doc, "因此，两天作出决定与“编辑初筛判定不适合送审”高度一致。没有拒稿信时不能 100% 锁定，但这是目前最强解释。即使论文创新更强，只要该政策仍在，重投同刊仍会面临相同结果。")

    add_subtitle(doc, "2.2 第二层原因：任务与组件已高度拥挤")
    add_table(
        doc,
        ["你的主张", "既有覆盖", "审计判断"],
        [
            ("ATT&CK 下一 parent-technique 预测", "BAN、CL-AP²、DeepOP、LiteATNet、Automated ATT&CK Technique Chaining 等已直接或高度邻近 [L3–L10]。", "任务本身不能作为创新点"),
            ("GRU/Transformer/Markov 序列建模", "Seq2Seq、DeepAG、DeepOP、ProAPT 已覆盖多种深序列预测 [L1–L2,L8,L10]。", "属于常规主干"),
            ("LLM 生成 reasoning 后语义增强", "AECR、SynthCTI、LLM-guided evidence mining 等已覆盖 LLM→语义/下游判别 [L14,L16–L18]。", "场景组合有价值，但不足以单独定新"),
            ("固定 α logit fusion", "标准 late fusion；M2GNN 已在非安全领域做 LLM/结构分支的逐样本动态 logit gate [L19]。", "算子创新低"),
            ("可解释下一步预测", "EFI/Nip in the Bud 已做预测与 ATT&CK 技术级解释 [L7]。", "需证明 faithfulness，不能只展示可读文本"),
        ],
        widths=[4.2, 8.4, 4.8],
        size=8.2,
    )

    add_subtitle(doc, "2.3 第三层原因：论文自己的结果没有形成强证据")
    add_bullet(doc, "表2中 Fusion 相对 GRU 的 Top-1 仅 0.5477−0.5444=0.0033，约等于 2,107 个测试前缀中的 7 个额外命中；Top-5 约增加 18 个命中。")
    add_bullet(doc, "二阶 Markov 的 Top-1=0.5562，高于所提 Fusion=0.5477；这削弱“新模型总体更强”的叙事。")
    add_bullet(doc, "Fusion 的 weighted-F1=0.5215，低于 GRU=0.5237；不能写成对所有指标一致提升。")
    add_bullet(doc, "表2与表3对相似方法/seed 给出明显不同的 MRR 和 Top-1；论文未解释运行、切分或 checkpoint 差异。")
    add_bullet(doc, "“低置信样本更受益”“排名改善多于恶化”是有意思的诊断，但尚未转化为可靠性机制或统计保证。")

    add_subtitle(doc, "2.4 P0 级实验溯源问题")
    add_table(
        doc,
        ["严重度", "直接证据", "影响", "必须动作"],
        [
            ("P0", "Multi-SeedAggregator.py:85–87 明示下方只是示例格式；:91–109 手填数组直接生成论文表3。", "表3没有自动实验溯源；不能作为论文证据。", "从逐 seed checkpoint/预测自动汇总，保存 manifest 与哈希"),
            ("P0", "固定 n=2,107 时，10 个 GRU/Fusion Top-1 手填原值中 7 个无法由整数命中数四舍五入得到。", "进一步表明当前数组不是有效的固定测试集逐样本统计。", "彻底删除现表3；重跑后报告整数命中数"),
            ("P0", "论文表2的 0.8695/0.6807/0.5477/0.8780/0.6861 在仓库源码/CSV/JSON/TXT 中无对应产物。", "表2无法复核。", "每个表格单元格追溯到 run_id 和 predictions.csv"),
            ("P0", "现存 rl_v2_test_predictions_top5.csv 重算为 Top-1 0.4893、Top-5 0.8306、MRR 0.6353。", "唯一完整单次产物不支持表2；但不能证明它就是论文 run。", "恢复原始 run 或重新运行，禁止用记忆/手抄数值"),
            ("P1", "核心 train/val/test CSV 的生成、过滤、父技术折叠和 split 脚本缺失。", "他人无法复现数据。", "提交 data_builder、split_manifest、版本和 seed"),
        ],
        widths=[1.3, 7.0, 4.3, 4.8],
        size=7.9,
    )
    add_callout(
        doc,
        "措辞边界",
        "可以下的结论是“表3由示例占位数组生成、当前无实验溯源，必须重算”；不能仅凭仓库推断作者主观伪造。科研修复的重点是建立自动、可复核的 evidence chain。",
        fill=LIGHT_GREY,
        accent="697782",
    )

    add_subtitle(doc, "2.5 论文—代码不一致")
    add_table(
        doc,
        ["项目", "论文描述", "代码事实", "影响"],
        [
            ("序列长度", "max prefix 50", "MAX_LEN=20；当前 CSV prefix_len 实际 1–19", "文字错误；当前数据暂未触发截断"),
            ("截断", "未强调方向", "训练 tokens[:20]，融合推理 seq[-20:]", "若未来长序列，训练/推理方向不一致"),
            ("Padding", "左填充、最后非 pad 状态", "右填充，直接使用 GRU 最终 hidden", "padding 参与状态演化，方法描述与实现不符"),
            ("损失", "标准 cross-entropy", "按当前真实标签排名乘 1/1.5/2 的 reward-weighted CE", "模型并非普通 GRU baseline"),
            ("checkpoint 选择", "统一流程不清", "GRU 按 Top-5/MRR；Transformer 按 MRR，loss 也不同", "baseline 比较不对称"),
            ("Macro-F1", "184 类", "sklearn 未传 labels=range(184)，test 只观察到 158 类", "不同模型可能使用不同分母"),
        ],
        widths=[2.3, 4.0, 6.1, 5.0],
        size=8.0,
    )

    add_subtitle(doc, "2.6 数据划分与统计独立性")
    add_table(
        doc,
        ["Split", "行数", "sequence_id", "SIM 根组", "出现标签"],
        [
            ("Train", "9,919", "552", "142", "184"),
            ("Validation", "2,102", "118", "75", "161"),
            ("Test", "2,107", "119", "73", "158"),
            ("合计", "14,128", "789", "162（唯一）", "184"),
        ],
        widths=[3.0, 2.6, 3.2, 3.3, 3.3],
        size=8.8,
    )
    add_bullet(doc, "精确 sequence_id 两两零重叠，sequence-level split 有证据支持。")
    add_bullet(doc, "若 SIM_xxx 是 campaign/scenario family，则 train-test 有 65 个根组重叠，占 73 个测试根组的 89%；41 个根组横跨三组。仓库没有真正 campaign 字段或 split manifest，因此 campaign-level 声明至少不可验证。")
    add_bullet(doc, "测试集 2,107 行来自 119 条序列、73 个根组；同一序列的嵌套前缀不独立。现有显著性脚本把 2,107 行直接用于 Wilcoxon/McNemar，需改为 sequence/campaign cluster bootstrap 或 paired permutation。")
    add_bullet(doc, "现存 GRU 产物从行级到序列等权、SIM 根组等权，Top-1 依次为 0.4893、0.4794、0.4697；评价口径可带来 1–2 个百分点差异，高于论文 0.33 个百分点的增益。")

    add_subtitle(doc, "2.7 消融、语义与部署证据")
    add_bullet(doc, "No-CoT 对齐脚本把 sequence_id 映射成字典；同一序列多个 prefix 只保留一条 reasoning，再广播到全部 prefix。这个控制实验不是逐 prefix 对齐，可能传播未来信息。")
    add_bullet(doc, "测试 reasoning 有 42/2,107 为空；Qwen 原始候选 Top-1 1.85%、Top-5 9.16%，且 250 行含标签表外 ID。下游提升实际来自“LLM 文本→BGE→监督 MLP”，不是 LLM zero-shot 预测。")
    add_bullet(doc, "reasoning 中直接出现 true label 的比例约 45.75%；进一步审计显示大部分是前缀中已出现标签的复述，约 4.6% 是新增标签候选。生成 prompt 未发现直接传 gold，但仍需要随机/交换/矛盾 rationale 控制，证明不是标签词捷径。")
    add_bullet(doc, "在线 BGE 编码使每样本成本约为 GRU 的 1,649 倍；当前成本表还不包含在线 LLM reasoning 生成，因此不能声称 realtime/online deployable。")
    add_bullet(doc, "现有校准结果是积极线索：Fusion ECE 12.14% 优于 GRU 13.46%，但仍明显失准，正好支持下一篇转向 calibration/selective prediction。")

    add_subtitle(doc, "2.8 写作与版面问题")
    add_bullet(doc, "第1–2页出现同一句话连续重复。")
    add_bullet(doc, "PDF 中大量绿色/红色链接框、引用框，呈现像编译调试稿。")
    add_bullet(doc, "第10页 Figure 3/Table 5、第11页表格字号过小，实际阅读困难。")
    add_bullet(doc, "多数参考文献缺期刊/会议、卷期、页码或 article number、DOI；与期刊要求的完整引用不符。")
    add_bullet(doc, "“8:1:1”与实际 70.21%/14.88%/14.91% 不一致；“consistently improves”与 Markov/weighted-F1 结果冲突。")

    add_review(
        doc,
        "任务1完成",
        [
            "主因与次因已分离：范围政策解释秒拒，创新/证据问题解释为何当前稿不宜直接转投。",
            "所有严重指控均降格为可验证的工程事实，不推断作者动机。",
            "结论同时覆盖论文方法、数值、数据、统计、消融、成本、写作和版式。",
        ],
        "最优决策不是润色后重投，而是先进行实验治理和贡献重构。若能提供正式拒稿信，应将本节第一层原因与编辑原文再对照一次。",
    )


def section_task2(doc: Document) -> None:
    add_section_title(doc, "3", "任务2｜不脱离现方向的新创新点")
    add_callout(
        doc,
        "主推荐",
        "把论文升级为“分歧感知的 conformal 选择性融合 + threat-actor/time/ATT&CK-version OOD 基准”。这条路线最大限度复用现有双分支代码，同时把创新从普通 late fusion 提升到可靠决策、有限样本覆盖、拒识与真实分布漂移。",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )

    add_subtitle(doc, "3.1 三条候选路线比较")
    add_table(
        doc,
        ["路线", "新颖性", "复用现代码", "主要风险", "建议"],
        [
            ("A. 分歧感知 conformal selective fusion + OOD", "中高：ATT&CK next-step 的覆盖/拒识/漂移仍有空白", "高：保留 GRU、语义 probe、fusion", "只换 MLP gate 会与既有动态融合近邻重合", "主论文"),
            ("B. ATT&CK 约束的反事实 near-miss listwise reranker", "中高：candidate-conditioned hard-negative + faithfulness 有空间", "中高：复用 top-k、LLM、SGLE-R 与重排代码", "容易变成知识图谱/LLM 组件堆叠", "第二优先或后续论文"),
            ("C. actor/time/version OOD benchmark", "中：评测协议价值高", "高：主要改数据与评估", "单独做可能只像 benchmark/analysis", "必须与 A 结合"),
        ],
        widths=[5.2, 4.0, 3.3, 3.5, 1.6],
        size=8.0,
    )

    add_subtitle(doc, "3.2 主路线 A+C：方法定义")
    add_number(doc, "双分支保留。序列分支输出 p_seq(y|x)；语义分支只看已观察前缀生成/编码语义证据，输出 p_sem(y|x)。")
    add_number(doc, "分支校准。分别使用 validation/calibration split 学习 temperature 或 vector scaling，不把 test 用于校准。")
    add_number(doc, "分歧特征。门控只使用可解释统计量：两支 entropy、top1–top2 margin、Jensen–Shannon divergence、候选 rank 差、类别频次、前缀稀有度、语义距离和 OOD score。")
    add_number(doc, "逐样本融合。学习 α(x)，但把 fixed α、stacking、rank fusion、plain MLP gate 作为强基线；门控本身不是核心创新。")
    add_number(doc, "Conformal 集合。按 tactic/频次层做 Mondrian calibration，输出候选集合 C(x)，报告 empirical coverage 与平均集合大小。")
    add_number(doc, "选择性预测。若集合过大、两分支冲突或 OOD score 超阈值，则 abstain/转人工，评价 risk–coverage 和 selective accuracy。")
    add_number(doc, "漂移评测。campaign/actor/report-source/time/ATT&CK version 分组，分别报告 in-domain、source-OOD、temporal-OOD、version-OOD。")

    add_subtitle(doc, "3.3 可写进论文的四项贡献")
    add_bullet(doc, "贡献1｜一个可审计的 ATT&CK next-technique 评测协议：group-disjoint、time-aware、version-aware，并公开 split manifest。")
    add_bullet(doc, "贡献2｜一个分歧感知的双分支选择性融合器：不把 gate 当创新终点，而是让冲突和长尾驱动预测集合/拒识。")
    add_bullet(doc, "贡献3｜parent-technique/tactic 条件化的 conformal prediction set，在可交换校准假设下给出有限样本覆盖目标；对分布漂移仅报告经验覆盖，不夸大保证。")
    add_bullet(doc, "贡献4｜对 LLM 语义证据的反事实审计：交换、打乱、删除、矛盾 rationale，验证 semantic branch 是否真正影响 near-miss 决策。")

    add_subtitle(doc, "3.4 研究问题与可证伪假设")
    add_table(
        doc,
        ["RQ", "问题", "通过门槛"],
        [
            ("RQ1", "分歧特征能否识别 fixed fusion 的失败样本？", "OOD/冲突子集 AUROC、error detection AUPRC 显著优于 max-softmax"),
            ("RQ2", "在给定 coverage 下，conformal selective fusion 是否降低 risk？", "跨 seed 与跨 split 的 risk–coverage 曲线优势；cluster bootstrap CI"),
            ("RQ3", "语义分支在何种 shift/长尾条件下真正有效？", "按 actor、版本、频次、前缀长度和 near-miss 分层"),
            ("RQ4", "reasoning 是有效证据还是标签词/长度捷径？", "随机/交换/矛盾 rationale 应产生可解释的性能下降"),
            ("RQ5", "覆盖保证在真实 shift 下如何退化？", "明确区分 exchangeable calibration guarantee 与 OOD empirical coverage"),
        ],
        widths=[1.4, 9.1, 6.2],
        size=8.2,
    )

    add_subtitle(doc, "3.5 备选路线 B：反事实 near-miss listwise reranker")
    add_body(doc, "如果你更希望保留“LLM reasoning”作为核心，可把语义分支改成 candidate-conditioned：GRU 先给 top-k；LLM 对每个候选分别生成结构化支持/反证；hard negatives 从同 tactic、语义邻近和 GRU top-k 近失类别抽取；以 pairwise/listwise loss 学习重排。ATT&CK tactic、platform、procedure dependency 只作约束/正则，不把它们包装成新知识图谱。")
    add_bullet(doc, "专门指标：near-miss subset MRR/NDCG、tactic-consistent error、ATT&CK graph distance、rare-class recall。")
    add_bullet(doc, "faithfulness：candidate swap、reason swap、删除关键证据、注入矛盾证据；若预测不变，则不能称 reasoning 驱动。")
    add_bullet(doc, "防泄漏：生成端只看前缀，不能看到 gold ID、完整序列或由 gold label 派生的描述；按 key 合并，禁止行号合并。")

    add_subtitle(doc, "3.6 明确不能再声称什么")
    add_bullet(doc, "“首次进行 ATT&CK 下一步预测”“全新 logit fusion”“因果解释”“multimodal”“end-to-end”。")
    add_bullet(doc, "“zero-day/unknown attack”“SOTA”“real-time”“显著降低分析师负担”，除非有对应协议、同分割基线、完整延迟/成本或用户研究。")
    add_bullet(doc, "固定 softmax entropy 不能等同于 uncertainty-aware；可读 rationale 不能等同于 faithful explanation。")

    add_review(
        doc,
        "任务2完成",
        [
            "主路线复用当前双分支与 parent-technique 任务，没有脱离研究方向。",
            "创新由方法、评测协议、可靠性指标和外部 OOD 四部分共同支撑，不依赖单一新模块。",
            "覆盖保证的适用条件已限定，避免在 distribution shift 下作错误理论承诺。",
        ],
        "这条路线有达到 Q2+ 审稿门槛的潜力，但不能保证录用；决定性因素是 P0 实验治理、公开 split 和外部 OOD 结果，而不是标题包装。",
    )


def section_task3(doc: Document) -> None:
    add_section_title(doc, "4", "任务3｜期刊网页与近邻文献重复性审计")
    add_callout(
        doc,
        "检索结论",
        "未发现与“GRU + LLM reasoning + BGE/MLP semantic logits + fixed α fusion”逐模块完全相同的论文；但任务重合高、组件创新弱。可辩护的新空白是 selective prediction、conformal set、拒识和 actor/time/version OOD，而不是 next-step 或 late fusion 本身。",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    add_subtitle(doc, "4.1 最接近的工作")
    add_table(
        doc,
        ["文献", "与当前稿重合", "可借鉴", "需要超越"],
        [
            ("DeepOP (2025) [L8]", "ATT&CK 具体下一技术、本体+深度序列", "因果/并行序列构造、强基线", "真实 OOD、校准/拒识、可审计切分"),
            ("Automated ATT&CK Technique Chaining (2025) [L9]", "语义依赖+数据驱动技术链", "ground-truth/可见性偏差分析", "有监督下一步与风险控制"),
            ("CL-AP² (2024) [L5]", "图、序列、具体技术预测", "attack portraying 与强模型对比", "分布漂移和选择性预测"),
            ("Nip in the Bud / EFI (2025) [L7]", "下一动作预测+技术级解释", "provenance/场景图与解释模板", "reasoning faithfulness 与覆盖风险"),
            ("BAN (2023) [L3]", "ATT&CK 下一技术概率预测", "概率图与防御映射", "校准的预测集合而非原始概率"),
            ("M2GNN (2026) [L19]", "LLM/结构双分支动态 logit gate（非安全）", "逐样本 α 与门控基线", "安全场景特有的 conformal/OOD 机制"),
            ("Conformal CPS (2026) [L20]", "网络安全中的 conformal 可靠性", "coverage、false alarm、不确定性", "ATT&CK 多类 next-step 与双分支冲突"),
            ("LLM evidence mining (2026) [L17]", "LLM 原型、对比证据、稀有技术", "hard negatives、证据学习", "candidate-conditioned 下一步与反事实 faithfulness"),
        ],
        widths=[4.5, 4.8, 4.0, 4.2],
        size=7.7,
    )

    add_subtitle(doc, "4.2 重合矩阵摘要")
    add_table(
        doc,
        ["维度", "当前工作", "文献拥挤度", "仍可创新"],
        [
            ("ATT&CK 下一技术", "核心任务", "高", "不再声明任务首创"),
            ("序列建模", "GRU", "高", "仅保留为基干"),
            ("LLM/文本语义", "reasoning→BGE→MLP", "中高", "以反事实信忠度验证形成差异"),
            ("late/logit fusion", "固定全局 α", "高", "分歧特征 + calibration + abstention"),
            ("near-miss", "有错误排名诊断", "中", "candidate-conditioned listwise 训练和专门指标"),
            ("UQ/OOD", "几乎缺失", "低到中", "conformal set、risk–coverage、actor/time/version shift"),
        ],
        widths=[3.2, 5.3, 3.0, 5.9],
        size=8.4,
    )

    add_subtitle(doc, "4.3 投稿期刊匹配")
    add_table(
        doc,
        ["候选", "官方范围匹配", "适合的论文版本", "风险"],
        [
            ("Journal of Network and Computer Applications", "明确接收 computer/network security applications [S2]", "完整 A+C；需外部攻击流程与网络应用价值", "不能只有模拟 CSV 和小幅 accuracy"),
            ("IEEE TDSC", "关注构建、建模与评估 dependable/secure systems 的基础、方法与机制 [S3]", "高标准 A+C；强调可靠性保证、攻击者/版本 shift", "方法与评测门槛最高"),
            ("Information Fusion", "面向 multi-source/multi-process fusion 的架构、算法与真实应用 [S4]", "融合方法必须有普适性、理论/消融，不只安全应用", "固定/普通 gate 远远不够"),
            ("International Journal of Information Security", "综合信息安全方法与应用；需提交前复核当年 JCR", "A+C 或 B，安全问题定义要强", "需要更完整真实数据"),
            ("Journal of Information Security and Applications", "安全应用与 ML 主题较匹配；作为较稳健备选", "完成 P0 后的 A+C 精简版", "当前分区每年变化，投稿前查最新 JCR"),
            ("Reliability Engineering & System Safety", "重视复杂系统、模型/参数不确定性与安全可靠性 [S5]", "仅当加入工业 CPS/OpTC/LANL 场景和实质可靠性问题", "纯 ATT&CK 文本序列可能 scope 不足"),
        ],
        widths=[4.2, 5.1, 5.0, 3.1],
        size=7.8,
    )
    add_body(doc, "期刊分区会随 JCR 年份、学科类别和学校口径变化。本报告只做研究匹配，不把历史 Q1/Q2 当成 2026 投稿保证；投稿前必须用本单位可访问的最新 JCR 再确认。")

    add_review(
        doc,
        "任务3完成",
        [
            "直接任务近邻、C&S 刊内近邻、LLM/fusion/UQ 方法近邻均已覆盖。",
            "没有发现 exact duplicate，但已明确当前版本的功能重复和组件重复。",
            "每条新路线都给出最近邻、可借鉴点与必须超越的边界。",
        ],
        "最危险的重复不是出现一篇完全相同论文，而是审稿人把当前方法归类为“标准 late fusion 在拥挤任务上的应用”。A+C 路线正是为改变这一归类。",
    )


def section_task4(doc: Document) -> None:
    add_section_title(doc, "5", "任务4｜实验数据集检索与可用性验证")
    add_callout(
        doc,
        "数据结论",
        "真正同时具有攻击顺序/图结构与 ATT&CK ID、可直接用于 next-step 的公开资源很少。核心应使用 CTID Attack Flow + Adversary Emulation Library；ATT&CK STIX 是标签/版本基础；TIE/TRAM/AnnoCTR 是无序或文本辅助；OpTC/LANL 是高成本遥测 OOD。",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )

    add_subtitle(doc, "5.1 可用性分级")
    add_table(
        doc,
        ["级别 / 数据", "实测规模与结构", "许可 / 验证", "建议用途"],
        [
            ("A｜CTID Attack Flow [D1]", "corpus 40 个 .afb、约 9.80 MB；有显式有向边、条件、分支和 technique_id/tactic_id", "Apache-2.0；GitHub Tree API、HEAD 和样例内容已验", "next-edge/next-technique 外部金标；保留分支，按事件/报告族分组"),
            ("A｜CTID AEL [D2]", "README 11 full + 12 micro plans；Tree 4,394 文件约 787.8 MB；机器 YAML 只覆盖部分 full", "Apache-2.0；仓库树、README、YAML 路径/大小已验", "第二独立序列域；actor/plan-family held-out OOD"),
            ("A｜ATT&CK STIX v19.1 [D3]", "index：Enterprise/Mobile 各 40 历史版本、ICS 26；当前 Enterprise JSON 约 53.3 MB", "ATT&CK 定制许可；index 完整解析，大文件 HEAD", "ID 规范化、deprecated/revoked、version-OOD；不是时间序列"),
            ("B｜CTID TIE [D4]", "完整 JSON 6,236 reports、47,586 非零 report-technique pairs、611 IDs", "Apache-2.0；5.26 MB JSON 完整解析", "候选/共现先验；无顺序，须排除外部测试来源"),
            ("B｜AnnoCTR [D5]", "400 CTI reports；仅 120 有完整 cyber tactic/technique 标注", "CC-BY-SA-4.0；Tree + 代表样本；仓库 archived", "文本→ATT&CK、vendor/time OOD；提及顺序非攻击顺序"),
            ("B｜TRAM2 [D6]", "single 5,089 rows/50 labels；multi 19,178 rows，其中 4,070 非空/50 labels", "Apache-2.0；两份 JSON 完整解析", "文本分类前端；必须按 doc_title 分组"),
            ("C｜OpTC [D7]", "官方约 1 TB、500 hosts；eCAR+Bro+red-team GT，有事件时间", "DARPA 公共领域/Distribution A；链接可达，未下载大文件", "强 telemetry OOD；需人工 ATT&CK 映射，非最小实验"),
            ("C｜LANL 58-day [D8]", "1,648,275,307 events、约 12 GB、58天；小型 redteam GT HEAD 已验", "LANL 权利放弃，近似 CC0；主数据未下载", "较轻量 telemetry case study；仍需人工映射"),
            ("D｜WAVE-27K [D9]", "论文称 27,801 CVE descriptions、27 techniques；无步骤顺序", "截至审计日未找到官方数据文件、代码或许可证", "不可作为可复现实验数据"),
        ],
        widths=[4.0, 5.3, 4.3, 3.8],
        size=7.35,
    )

    add_subtitle(doc, "5.2 最小可复现实验组合")
    add_number(doc, "以 ATT&CK STIX v15 固定训练标签快照，建立 parent/sub-technique、revoked/deprecated 与别名映射；以 v19.1 构造新增/变化标签的 version-OOD。")
    add_number(doc, "用清洗后的 TIE 只训练 candidate/co-occurrence prior；移除来源为 Attack Flow 的 29 条、AEL 的 10 条，并继续按 report/actor/campaign 去近重复。")
    add_number(doc, "将 Attack Flow 的有向边转为“已观察节点/前缀 → 下一 technique 集合”；分支场景使用多标签 next-set，不任意线性化成唯一顺序。")
    add_number(doc, "将 AEL 的完整计划步骤转换为第二序列域；按 actor/plan family 留出，作为 source-OOD。")
    add_number(doc, "AnnoCTR/TRAM2 只验证 CTI 文本到 ATT&CK 技术的前端，不能把文档提及顺序冒充攻击进程。")
    add_number(doc, "资源允许时，选择 LANL red-team 小范围人工映射 case study；OpTC 作为扩展实验，而非首轮必做。")

    add_subtitle(doc, "5.3 关键泄漏与复现风险")
    add_bullet(doc, "TIE 完整 JSON 的 origin_of_data 实测：OpenCTI 6,025、TRAM 149、Attack Flow 29、ATT&CK Campaigns 22、AEL 10、VECTR 1。若 Attack Flow/AEL 用作外部测试，必须先移除对应来源。")
    add_bullet(doc, "TIE README 的 43,899 observations 与当前文件的 47,586 非零 pair/49,994 频次总数不一致；应固定下载日期、文件哈希和统计口径。")
    add_bullet(doc, "AEL 是仿真剧本，不代表真实攻击频率；Attack Flow 通常没有真实执行时间；不能用它们证明实时预测。")
    add_bullet(doc, "OpTC/LANL ground truth 并非逐事件 ATT&CK 标签；映射需双人复核、记录证据与不确定标签，避免事后主观贴标。")

    add_subtitle(doc, "5.4 验证深度说明")
    add_table(
        doc,
        ["验证强度", "已覆盖"],
        [
            ("整文件下载并解析", "ATT&CK index.json；TIE combined_dataset_full_frequency.json；TRAM2 single/multi JSON"),
            ("目录 API + README + 小样本/HEAD", "ATT&CK Enterprise、Attack Flow、AEL、AnnoCTR、LANL red-team"),
            ("仅链接/官方元数据", "OpTC 与 LANL 大型主数据；未宣称本地完整验证"),
            ("只有论文元数据", "WAVE-27K；判定为当前不可复现"),
        ],
        widths=[4.4, 12.1],
        size=8.6,
    )

    add_review(
        doc,
        "任务4完成",
        [
            "每个数据资源均给出官方入口、许可、结构、顺序性、ATT&CK ID 和验证深度。",
            "已区分可直接 next-step、无序先验、文本辅助、遥测 OOD 与不可复现资源。",
            "已发现并设计规避 TIE→Attack Flow/AEL 的来源泄漏。",
        ],
        "“链接能打开”不等于“数据集可用”。本报告只把已经解析或至少验证目录/样例/许可的资源列为 A/B；大型遥测明确标为 C。",
    )


def section_roadmap(doc: Document) -> None:
    add_section_title(doc, "6", "实施路线｜从当前仓库到可投稿稿件")
    add_subtitle(doc, "6.1 四阶段路线")
    add_table(
        doc,
        ["阶段", "工作", "验收门槛", "建议周期"],
        [
            ("P0 科研治理", "重建数据构造/split；修复 padding、hidden、No-CoT；统一 metrics；自动 run manifest", "论文每个数字可从 predictions 重算；表3不再含手填数组", "2–3 周"),
            ("P1 数据与基线", "接入 Attack Flow/AEL/STIX；TIE 清洗；Markov/n-gram/GRU/LSTM/Transformer/固定融合/动态 gate", "actor/campaign/report/time/version split 固化并公开", "3–4 周"),
            ("P2 新方法", "校准、分歧 gate、Mondrian conformal、abstention、OOD score、rationale corruption", "覆盖/集合大小/risk-coverage/校准 + 排名指标全部报告", "4–6 周"),
            ("P3 论文与复现", "五种子、cluster CI、成本、错误分析、ablation、artifact package、重写论文", "一键生成表图；外部 OOD 至少两域；无未解释数字", "3–4 周"),
        ],
        widths=[2.6, 7.1, 5.5, 2.4],
        size=8.0,
    )

    add_subtitle(doc, "6.2 实验矩阵最低门槛")
    add_table(
        doc,
        ["类别", "必须包含"],
        [
            ("基线", "Markov/n-gram、GRU、LSTM、Transformer、固定 α、learned gate、rank fusion；有条件加入 Bayesian/BAN-style"),
            ("数据切分", "random/sequence 仅辅助；主结果为 campaign/actor/report-source/time/version-disjoint"),
            ("语义消融", "no rationale、empty、random、shuffle、swap、contradictory、candidate-description-only、oracle gate"),
            ("准确/排序", "Top-k/Recall@k、MRR、NDCG、固定 184 labels 的 macro-F1、rare-class recall"),
            ("可靠性", "ECE、classwise ECE、Brier、coverage、set size、risk–coverage、selective accuracy"),
            ("统计", "五种子；sequence/campaign cluster bootstrap CI；paired cluster permutation；效应量"),
            ("部署", "GRU、BGE、probe、fusion、在线 LLM 分项延迟/吞吐/成本；缓存与失败降级"),
        ],
        widths=[3.3, 13.1],
        size=8.5,
    )

    add_subtitle(doc, "6.3 推荐仓库产物结构")
    add_code_line(doc, "artifacts/<run_id>/config.yaml")
    add_code_line(doc, "artifacts/<run_id>/data_manifest.json        # source URL/hash/version/split IDs")
    add_code_line(doc, "artifacts/<run_id>/predictions_test.parquet # sample_id, group_id, logits, rank, abstain")
    add_code_line(doc, "artifacts/<run_id>/metrics.json             # 自动从 predictions 计算")
    add_code_line(doc, "artifacts/<run_id>/environment.lock         # Python/CUDA/model revisions")
    add_code_line(doc, "reports/tables/*.csv                         # 自动聚合，不允许手填")
    add_code_line(doc, "reports/figures/*.pdf                        # 由脚本生成")

    add_subtitle(doc, "6.4 摘要/贡献写法模板")
    add_callout(
        doc,
        "推荐叙事",
        "现有 ATT&CK next-technique 模型在同分布准确率上表现良好，但在威胁主体、报告来源和知识库版本变化时缺少可靠性控制。我们提出一个分歧感知的双分支选择性融合框架，输出校准预测集合并在高风险样本上拒识；同时发布可审计的 group/time/version OOD 协议。",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )
    add_body(doc, "不要再把 0.3–1.7 个百分点的总体提升放在摘要中心。摘要中心应是：问题缺口（shift 下不可靠）、机制（disagreement + conformal + abstention）、协议（actor/time/version OOD）和实际风险收益（coverage/risk/外部域）。")

    add_subtitle(doc, "6.5 投稿决策")
    add_bullet(doc, "当前版本：不投。先完成 P0。")
    add_bullet(doc, "完成 P0+P1，但新方法未完成：可作为技术报告/预印本的复现修正版，不宜声称 Q2+ 方法创新。")
    add_bullet(doc, "完成 A+C、两域外部 OOD、cluster 统计和 artifact：优先评估 JNCA；方法学足够强再冲 TDSC/Information Fusion。")
    add_bullet(doc, "投稿前再次核验目标期刊最新 scope 和 JCR 分区；特别是 C&S 的 AI/ML 政策是否仍在。")

    add_review(
        doc,
        "整体方案可执行",
        [
            "每个创新都映射到具体数据、基线、消融、指标和统计检验。",
            "优先修复证据链，再增加模型；不会用新模块掩盖旧实验不可复现。",
            "投稿目标按方法强度和真实外部验证分层，不承诺分区或录用。",
        ],
        "如果只能做一件事，先重建自动可追溯实验；如果只能做一个新创新，做 conformal selective fusion + actor/time/version OOD，而不是再换一个序列编码器。",
    )


def section_appendix_code(doc: Document) -> None:
    add_section_title(doc, "附录A", "源码与产物证据索引")
    add_table(
        doc,
        ["证据主题", "本地路径 / 行号"],
        [
            ("核心数据", r"project\data\sim_{train,val,test}_parent_min3.csv"),
            ("标签表", r"project\data\rl_label_vocab.csv"),
            ("GRU 实现", r"project\rl\train_rl_baseline_v2.py:43,60–64,121–168,258–267"),
            ("GRU 多种子", r"project\rl\train_rl_multiseed.py:19–34,94–110,127–217"),
            ("LLM 生成", r"project\llm\run_llm_kg_context_test_pipeline.py:13–20,42–46,70–135"),
            ("语义探针", r"project\llm\train_llm_multiseed.py:39–70,79–162"),
            ("全局融合", r"project\fusion\run_final_fusion.py:40–72,101–153,178–233"),
            ("表3手填来源", r"project\data\Multi-SeedAggregator.py:83–109,121–134"),
            ("表3 CSV", r"project\data\multiseed_summary_table.csv"),
            ("单次预测", r"project\rl\rl_v2_test_predictions_top5.csv"),
            ("显著性脚本", r"project\data\stat_test_sample_level.py:53–69,87–232"),
            ("No-CoT 对齐", r"project\Ablation experiment\hotfix_align.py:18–27"),
            ("校准", r"project\data\calibration_summary.csv"),
            ("成本", r"project\data\cost_breakdown_system.csv; cost_breakdown_components.csv"),
            ("外部 CTID", r"project\rl\all_ctid_eval\rl_all_ctid_summary.csv"),
            ("micro-state", r"project\data_v2\micro_state\state_vs_microstate_knn_metrics.csv"),
            ("56 样本动态融合", r"project\data_v2\micro_state\dynamic_rl_llm_fusion_56_metrics.txt"),
        ],
        widths=[5.0, 11.5],
        size=8.0,
    )

    add_subtitle(doc, "A.1 已核实的关键重算")
    add_bullet(doc, "行数：9,919 + 2,102 + 2,107 = 14,128；标签表 184 个唯一 parent-techniques。")
    add_bullet(doc, "现存单次 GRU 逐样本预测：Top-1 0.489321、Top-5 0.830565、MRR 0.635258。")
    add_bullet(doc, "外部 CTID 旧 GRU：10 个组织、281 样本加权汇总约 Top-1 2.85%、Top-5 9.96%、MRR 0.077。该结果不能直接代表论文 Fusion，但显示明显域差距。")
    add_bullet(doc, "micro-state KNN：micro_state Top-1 0.5125，纯 state 0；56 样本中 LLM-only Top-1 0.3571，旧 RL 0.0357，动态融合未改善。")


def section_references(doc: Document) -> None:
    add_section_title(doc, "附录B", "外部来源与可点击链接")

    sources = [
        ("S1", "Computers & Security — Guide for Authors（含当前 AI/ML moratorium）", "https://www.sciencedirect.com/journal/computers-and-security/publish/guide-for-authors"),
        ("S2", "Journal of Network and Computer Applications — Aims & Scope", "https://www.sciencedirect.com/journal/journal-of-network-and-computer-applications"),
        ("S3", "IEEE Transactions on Dependable and Secure Computing — Scope overview", "https://cybersecurity.ieee.org/read/"),
        ("S4", "Information Fusion — Aims & Scope", "https://www.sciencedirect.com/journal/information-fusion"),
        ("S5", "Reliability Engineering & System Safety — Aims & Scope", "https://www.sciencedirect.com/journal/reliability-engineering-and-system-safety"),
    ]
    add_subtitle(doc, "B.1 期刊与政策")
    for key, title, url in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"[{key}] {title}. ")
        set_run_font(r, size=8.8, color=DARK)
        add_hyperlink(p, url, url)

    literature = [
        ("L1", "Detecting multi-stage attacks using sequence-to-sequence model", "https://doi.org/10.1016/j.cose.2021.102203"),
        ("L2", "DeepAG: Attack Graph Construction and Threats Prediction with Bi-directional Deep Learning", "https://doi.org/10.1109/TDSC.2022.3143551"),
        ("L3", "BAN: Predicting APT Attack Based on Bayesian Network With MITRE ATT&CK Framework", "https://doi.org/10.1109/ACCESS.2023.3306593"),
        ("L4", "A Causal Graph-Based Approach for APT Predictive Analytics", "https://doi.org/10.3390/electronics12081849"),
        ("L5", "CL-AP²: A composite learning approach to attack prediction via attack portraying", "https://doi.org/10.1016/j.jnca.2024.103963"),
        ("L6", "LiteATNet: Predicting APT Attack Using Transformer Model With MITRE ATT&CK Framework", "https://doi.org/10.1109/SWC62898.2024.00065"),
        ("L7", "Nip in the Bud: Forecasting and Interpreting Post-exploitation Attacks", "https://doi.org/10.1109/TDSC.2024.3444781"),
        ("L8", "DeepOP: A Hybrid Framework for MITRE ATT&CK Sequence Prediction", "https://doi.org/10.3390/electronics14020257"),
        ("L9", "Automated ATT&CK Technique Chaining", "https://doi.org/10.1145/3696013"),
        ("L10", "ProAPT: Projection of APTs with Deep Reinforcement Learning", "https://doi.org/10.22042/isecure.2024.428569.1052"),
        ("L11", "Detecting APT attacks using an attack intent-driven and sequence-based learning approach", "https://doi.org/10.1016/j.cose.2024.103748"),
        ("L12", "Automated discovery and mapping ATT&CK tactics and techniques for unstructured CTI", "https://doi.org/10.1016/j.cose.2024.103815"),
        ("L13", "Malware2ATT&CK: A sophisticated model for mapping malware to ATT&CK techniques", "https://doi.org/10.1016/j.cose.2024.103772"),
        ("L14", "AECR: Automatic attack technique intelligence extraction based on fine-tuned LLM", "https://doi.org/10.1016/j.cose.2024.104213"),
        ("L15", "Investigating co-occurrences of MITRE ATT&CK techniques", "https://doi.org/10.1016/j.cose.2026.105042"),
        ("L16", "SynthCTI: LLM-driven synthetic CTI generation to enhance MITRE technique mapping", "https://doi.org/10.1016/j.future.2025.108232"),
        ("L17", "LLM-guided contrastive evidence mining for explainable CTI classification", "https://doi.org/10.1016/j.isci.2026.116466"),
        ("L18", "Proactive Cyber Defense: A Real-Time CTI Framework with ATT&CK–D3FEND Mapping", "https://doi.org/10.3390/systems14050575"),
        ("L19", "Multimodal graph neural network with large language models for node and link prediction", "https://doi.org/10.3389/frai.2026.1758852"),
        ("L20", "Conformal machine learning for reliable anomaly detection in industrial CPS", "https://doi.org/10.1016/j.ress.2026.112417"),
    ]
    add_subtitle(doc, "B.2 近邻文献")
    for key, title, url in literature:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"[{key}] {title}. ")
        set_run_font(r, size=8.4, color=DARK)
        add_hyperlink(p, url, url)

    datasets = [
        ("D1", "CTID Attack Flow", "https://github.com/center-for-threat-informed-defense/attack-flow"),
        ("D2", "CTID Adversary Emulation Library", "https://github.com/center-for-threat-informed-defense/adversary_emulation_library"),
        ("D3", "MITRE ATT&CK STIX Data", "https://github.com/mitre-attack/attack-stix-data"),
        ("D4", "CTID Technique Inference Engine", "https://github.com/center-for-threat-informed-defense/technique-inference-engine"),
        ("D5", "Bosch AnnoCTR", "https://github.com/boschresearch/anno-ctr-lrec-coling-2024"),
        ("D6", "CTID TRAM", "https://github.com/center-for-threat-informed-defense/tram"),
        ("D7", "DARPA OpTC Data", "https://github.com/FiveDirections/OpTC-data"),
        ("D8", "LANL Comprehensive 58-day Cyber Security Data", "https://lanl.ma.ic.ac.uk/data/cyber1/"),
        ("D9", "WAVE-27K paper", "https://aclanthology.org/2024.nlpaics-1.14/"),
    ]
    add_subtitle(doc, "B.3 数据集")
    for key, title, url in datasets:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"[{key}] {title}. ")
        set_run_font(r, size=8.6, color=DARK)
        add_hyperlink(p, url, url)


def section_final_review(doc: Document) -> None:
    add_section_title(doc, "附录C", "最终 Review 清单")
    add_table(
        doc,
        ["任务", "完成证据", "最终状态"],
        [
            ("1. 创新不足/秒拒", "期刊官方政策 + 论文结果 + 源码/产物/版式联合归因", "完成"),
            ("2. 新创新点", "主路线 A+C、备选 B、RQ/方法/门槛/风险均定义", "完成"),
            ("3. 查重与借鉴", "20 篇近邻、重合矩阵、目标期刊范围与差异边界", "完成"),
            ("4. 数据集", "9 类资源分级；关键 JSON 完整解析；许可/顺序/ATT&CK ID/泄漏均核验", "完成"),
            ("5. 文档", "章节化报告、可点击外链、代码证据索引、阶段 Review", "完成"),
        ],
        widths=[4.1, 10.5, 2.0],
        size=8.6,
    )
    add_callout(
        doc,
        "最终决策",
        "不要对现稿做“小修式转投”。先把当前工作降级为可复现基线，清空不可追溯表格，再以“分歧感知 conformal selective fusion + actor/time/version OOD”为新主线重做实验。这样既保留你已经完成的序列/语义代码，又能形成审稿人可清晰识别的可靠性贡献。",
        fill=LIGHT_GREEN,
        accent=GREEN,
    )
    add_subtitle(doc, "提交前 12 个硬门槛")
    gates = [
        "拒稿信与目标期刊 scope 已再次核对。",
        "数据构造脚本、ATT&CK 版本、source URL/hash 已固化。",
        "campaign/actor/report/time split manifest 已发布。",
        "所有表图从逐样本预测自动生成，无手填实验数值。",
        "训练/验证/推理 padding、truncation、hidden state 完全一致。",
        "No-CoT/random/swap/contradictory rationale 按 sample key 对齐。",
        "Macro-F1 使用固定 184 labels；Top-k 同时报整数命中数。",
        "统计以 sequence/campaign 为 cluster，并给 CI/效应量。",
        "Attack Flow/AEL 外部测试已从 TIE 训练来源移除。",
        "conformal 保证的 exchangeability 条件与 OOD empirical coverage 已分开。",
        "在线 LLM、BGE、probe、fusion 成本全部计入。",
        "摘要不含 first/SOTA/causal/realtime/zero-day 等无证据措辞。",
    ]
    for item in gates:
        add_bullet(doc, "□ " + item)


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    configure_sections(doc)
    add_cover(doc)
    add_manual_contents(doc)
    section_executive(doc)
    section_scope(doc)
    section_task1(doc)
    section_task2(doc)
    section_task3(doc)
    section_task4(doc)
    section_roadmap(doc)
    section_appendix_code(doc)
    section_references(doc)
    section_final_review(doc)

    core = doc.core_properties
    core.title = "ZDS 论文拒稿复盘与 SCI Q2+ 重构路线（证据审计版）"
    core.subject = "ATT&CK next-technique prediction paper and code audit"
    core.author = "Codex"
    core.keywords = "MITRE ATT&CK, next-step prediction, conformal prediction, OOD, reproducibility"
    core.comments = "Generated from a read-only audit of the supplied paper, source code, artifacts, official journal pages, literature and datasets."
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build())
