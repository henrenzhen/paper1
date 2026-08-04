from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image


ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "deliverables" / "ZDS论文拒稿复盘与SCI_Q2+重构路线_证据审计版.docx"
PDF = ROOT / "tmp" / "docx_render" / "zds_report" / "zds_report_fallback.pdf"
PAGES = ROOT / "tmp" / "docx_render" / "zds_report" / "pages_v2"


def main() -> None:
    failures: list[str] = []
    if not DOCX.exists() or DOCX.stat().st_size < 50_000:
        failures.append("DOCX missing or unexpectedly small")

    with zipfile.ZipFile(DOCX) as archive:
        bad_member = archive.testzip()
        if bad_member:
            failures.append(f"corrupt ZIP member: {bad_member}")
        names = set(archive.namelist())
        for required in ("[Content_Types].xml", "word/document.xml", "word/styles.xml", "word/_rels/document.xml.rels"):
            if required not in names:
                failures.append(f"missing DOCX member: {required}")

    document = Document(DOCX)
    texts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    full_text = "\n".join(texts)

    required_phrases = [
        "任务1｜为什么创新不足，为什么会被秒拒",
        "任务2｜不脱离现方向的新创新点",
        "任务3｜期刊网页与近邻文献重复性审计",
        "任务4｜实验数据集检索与可用性验证",
        "阶段 Review · 任务1完成",
        "阶段 Review · 任务2完成",
        "阶段 Review · 任务3完成",
        "阶段 Review · 任务4完成",
        "附录C  最终 Review 清单",
        "Computers & Security",
        "conformal selective fusion",
        "CTID Attack Flow",
    ]
    for phrase in required_phrases:
        if phrase not in full_text:
            failures.append(f"missing required phrase: {phrase}")

    forbidden = ["TODO", "TBD", "PLACEHOLDER", "Lorem ipsum", "\ufffd"]
    for token in forbidden:
        if token in full_text:
            failures.append(f"forbidden token present: {token!r}")

    hyperlinks = []
    for rel in document.part.rels.values():
        if rel.reltype == RT.HYPERLINK:
            hyperlinks.append(rel.target_ref)
            if not rel.target_ref.startswith(("https://", "http://")):
                failures.append(f"invalid external hyperlink: {rel.target_ref}")
    if len(hyperlinks) < 30:
        failures.append(f"too few hyperlinks: {len(hyperlinks)}")

    page_files = sorted(PAGES.glob("page-*.png"))
    if len(page_files) != 13:
        failures.append(f"expected 13 QA pages, found {len(page_files)}")
    page_sizes = []
    for page in page_files:
        with Image.open(page) as image:
            page_sizes.append(image.size)
            if image.width < 900 or image.height < 1200:
                failures.append(f"unexpectedly small QA page: {page.name} {image.size}")
    if len(set(page_sizes)) > 1:
        failures.append(f"inconsistent QA page sizes: {sorted(set(page_sizes))}")

    result = {
        "docx": str(DOCX),
        "docx_bytes": DOCX.stat().st_size,
        "sha256": hashlib.sha256(DOCX.read_bytes()).hexdigest(),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "hyperlinks": len(hyperlinks),
        "qa_pdf_bytes": PDF.stat().st_size if PDF.exists() else 0,
        "qa_pages": len(page_files),
        "qa_page_size": page_sizes[0] if page_sizes else None,
        "failures": failures,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    if failures:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
