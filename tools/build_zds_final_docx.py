from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN = ROOT / "docs" / "research" / "2026-07-30-zds-desk-reject-and-dr-vom-final-report.md"
RESULTS = ROOT / "project" / "experiments" / "gsad" / "results" / "external" / "dr_vom_full_lodo_final_seed20260730"
ASSETS = ROOT / "deliverables" / "assets"
OUTPUT = ROOT / "deliverables" / "ZDS论文拒稿复盘与DR-VOM实验报告_最终版.docx"
CHART = ASSETS / "dr_vom_lodo_performance.png"

NAVY = "1F4D78"
BLUE = "2E74B5"
INK = "1B2430"
MUTED = "5B6573"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"
GRID = "B9C2CC"
RED = "9B1C1C"
GOLD = "7A5A00"
GREEN = "2E6B4F"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_borders(table, color=GRID, size=4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int], indent=120) -> None:
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=None, bold=None, italic=None, color=None, mono=False) -> None:
    latin = "Consolas" if mono else "Calibri"
    east = "Microsoft YaHei"
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_border(paragraph, side: str, color: str, size=12, space=6) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    node = pbdr.find(qn(f"w:{side}"))
    if node is None:
        node = OxmlElement(f"w:{side}")
        pbdr.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:space"), str(space))
    node.set(qn("w:color"), color)


def set_para_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rpr.append(fonts)
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\[[^\]]+\]\(https?://[^)]+\)|\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def add_inline(paragraph, text: str, size=11, color=INK) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=max(size - 0.5, 7.5), color=NAVY, mono=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=size, color=color, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size=size, color=color)


def add_page_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def new_numbering_instance(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        for style in abstract.iter(qn("w:pStyle")):
            if style.get(qn("w:val")) == "ListNumber":
                abstract_id = int(abstract.get(qn("w:abstractNumId")))
                break
        if abstract_id is not None:
            break
    if abstract_id is None:
        raise RuntimeError("List Number abstract numbering definition not found")
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def set_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    add_inline(p, "ZDS RESEARCH AUDIT  /  FINAL EVIDENCE REVIEW", size=8.5, color=MUTED)
    set_para_border(p, "bottom", "D7DBE2", size=4, space=3)
    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.add_run("Page ")
    add_page_field(fp, "PAGE")
    fp.add_run(" of ")
    add_page_field(fp, "NUMPAGES")
    for run in fp.runs:
        set_font(run, size=9, color=MUTED)
    section.first_page_footer.paragraphs[0].text = ""


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(116)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("RESEARCH AUDIT  ·  FINAL EVIDENCE REVIEW")
    set_font(run, size=10.5, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("ZDS 论文拒稿复盘与\nDR-VOM 最终实验报告")
    set_font(run, size=28, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    run = subtitle.add_run("从 LLM 语义后融合转向来源域 / 事件根平衡的 ATT&CK 下一技术预测")
    set_font(run, size=13.5, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(28)
    set_para_border(rule, "bottom", BLUE, size=12, space=1)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(9)
    add_inline(meta, "证据截止：2026-07-30  |  论文：Enhancing ATT&CK Parent-Technique Next-Step Prediction", size=10, color=MUTED)

    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.paragraph_format.space_before = Pt(18)
    status.paragraph_format.space_after = Pt(0)
    run = status.add_run("结论：方向有效，但尚未达到 SCI 二区投稿就绪门槛")
    set_font(run, size=11.5, bold=True, color=RED)

    doc.add_page_break()


def chart_data() -> tuple[list[str], list[float], list[tuple[float, float]], list[float], list[tuple[float, float]]]:
    with (RESULTS / "domain_metrics.csv").open(encoding="utf-8", newline="") as handle:
        metrics = list(csv.DictReader(handle))
    with (RESULTS / "bootstrap_intervals.csv").open(encoding="utf-8", newline="") as handle:
        intervals = list(csv.DictReader(handle))
    lookup = {(row["heldout_domain"], row["metric"]): row for row in intervals}
    labels = ["SIM", "CTID", "Attack Flow", "Stockpile"]
    keys = ["sim", "ctid", "attack_flow", "stockpile"]
    top = [float(row["top1_gain_pp"]) for row in metrics]
    mrr = [float(row["mrr_gain"]) for row in metrics]
    top_ci = [(float(lookup[(key, "top1_gain_pp")]["lower"]), float(lookup[(key, "top1_gain_pp")]["upper"])) for key in keys]
    mrr_ci = [(float(lookup[(key, "mrr_gain")]["lower"]), float(lookup[(key, "mrr_gain")]["upper"])) for key in keys]
    return labels, top, top_ci, mrr, mrr_ci


def make_chart() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    labels, top, top_ci, mrr, mrr_ci = chart_data()
    width, height = 1656, 684
    canvas = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(canvas)
    font_dir = Path("C:/Windows/Fonts")
    regular_path = font_dir / "arial.ttf"
    bold_path = font_dir / "arialbd.ttf"
    regular = ImageFont.truetype(str(regular_path), 24)
    small = ImageFont.truetype(str(regular_path), 20)
    bold = ImageFont.truetype(str(bold_path), 27)
    title_font = ImageFont.truetype(str(bold_path), 32)

    title = "DR-VOM leave-one-source-domain-out performance"
    box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((width - (box[2] - box[0])) / 2, 25), title, font=title_font, fill="#1F4D78")

    colors = ["#2E74B5", "#4E87B8", "#6A9CC7", "#8CB2D3"]
    panels = (
        ((70, 105, 805, 600), top, top_ci, "Top-1", "Gain (percentage points)", (-6.0, 16.0), 5.0),
        ((885, 105, 1620, 600), mrr, mrr_ci, "MRR", "Gain", (-0.05, 0.16), 0.05),
    )
    for (left, top_px, right, bottom), values, cis, panel_title, y_label, limits, tick_step in panels:
        plot_left, plot_top = left + 95, top_px + 55
        plot_right, plot_bottom = right - 20, bottom - 80
        y_min, y_max = limits

        def y_pos(value: float) -> float:
            return plot_bottom - (value - y_min) / (y_max - y_min) * (plot_bottom - plot_top)

        panel_box = draw.textbbox((0, 0), panel_title, font=bold)
        draw.text(((left + right - (panel_box[2] - panel_box[0])) / 2, top_px), panel_title, font=bold, fill="#1F4D78")
        tick = y_min
        while tick <= y_max + 1e-9:
            y = y_pos(tick)
            draw.line((plot_left, y, plot_right, y), fill="#E5E9EE", width=2)
            label = f"{tick:.2f}" if abs(tick_step) < 1 else f"{tick:.0f}"
            tick_box = draw.textbbox((0, 0), label, font=small)
            draw.text((plot_left - 12 - (tick_box[2] - tick_box[0]), y - 10), label, font=small, fill="#5B6573")
            tick += tick_step
        zero_y = y_pos(0.0)
        draw.line((plot_left, zero_y, plot_right, zero_y), fill="#27313C", width=3)
        draw.line((plot_left, plot_top, plot_left, plot_bottom), fill="#27313C", width=2)
        n = len(labels)
        slot = (plot_right - plot_left) / n
        bar_width = slot * 0.55
        for idx, (label, value, ci) in enumerate(zip(labels, values, cis)):
            center = plot_left + slot * (idx + 0.5)
            value_y = y_pos(value)
            draw.rectangle((center - bar_width / 2, min(value_y, zero_y), center + bar_width / 2, max(value_y, zero_y)), fill=colors[idx])
            low_y, high_y = y_pos(ci[0]), y_pos(ci[1])
            draw.line((center, high_y, center, low_y), fill="#27313C", width=3)
            draw.line((center - 10, high_y, center + 10, high_y), fill="#27313C", width=3)
            draw.line((center - 10, low_y, center + 10, low_y), fill="#27313C", width=3)
            label_box = draw.textbbox((0, 0), label, font=small)
            draw.text((center - (label_box[2] - label_box[0]) / 2, plot_bottom + 15), label, font=small, fill="#27313C")
        y_box = draw.textbbox((0, 0), y_label, font=regular)
        y_image = Image.new("RGBA", (y_box[2] - y_box[0] + 10, y_box[3] - y_box[1] + 10), (255, 255, 255, 0))
        y_draw = ImageDraw.Draw(y_image)
        y_draw.text((5, 5 - y_box[1]), y_label, font=regular, fill="#27313C")
        rotated = y_image.rotate(90, expand=True)
        canvas.paste(rotated, (left + 5, int((plot_top + plot_bottom - rotated.height) / 2)), rotated)
    canvas.save(CHART, format="PNG", optimize=True)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows


def table_widths(rows: list[list[str]]) -> list[int]:
    n = len(rows[0])
    header = rows[0]
    if n == 7 and "留出来源" in header[0]:
        return [1150, 1030, 1100, 1100, 2200, 1800, 980]
    if n == 5:
        return [1450, 2450, 1550, 1550, 2360]
    if n == 4:
        return [1850, 2500, 2300, 2710]
    if n == 3:
        return [1800, 2500, 5060]
    if n == 2:
        return [2700, 6660]
    base = 9360 // n
    widths = [base] * n
    widths[-1] += 9360 - sum(widths)
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = table_widths(rows)
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        set_row_cant_split(row)
    compact = len(rows[0]) >= 5
    body_size = 7.7 if len(rows[0]) == 7 else (8.4 if compact else 9.0)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            is_numeric = bool(re.search(r"\d", value)) and len(value) < 34
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if (r_idx == 0 or is_numeric) else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(paragraph, value, size=body_size, color=INK)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.13)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    set_para_shading(p, CALLOUT)
    set_para_border(p, "left", BLUE, size=16, space=7)
    add_inline(p, text, size=10.5, color=INK)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.05
    set_para_shading(p, "F7F8FA")
    set_para_border(p, "left", "AAB6C3", size=10, space=5)
    for idx, line in enumerate(lines):
        run = p.add_run(("\n" if idx else "") + line)
        set_font(run, size=8.3, color="263442", mono=True)


def add_body_paragraph(doc: Document, text: str, style=None, size=11):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.widow_control = True
    add_inline(p, text.rstrip("  "), size=size, color=INK)
    return p


def add_figure(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    inline_shape = run.add_picture(str(CHART), width=Inches(6.15))
    inline_shape._inline.docPr.set(
        "descr",
        "DR-VOM 在 SIM、CTID、Attack Flow 和 Stockpile 四个留一来源域上的 Top-1 与 MRR 增益及 95% root-cluster bootstrap 置信区间",
    )
    inline_shape._inline.docPr.set("title", "DR-VOM 四来源域性能图")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    run = caption.add_run("图 1  四个留一来源上的 Top-1 与 MRR 增益（误差线为 root-cluster bootstrap 95% CI）")
    set_font(run, size=9, italic=True, color=MUTED)


def add_markdown_body(doc: Document) -> None:
    lines = MARKDOWN.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("> **最终判断"))
    i = start
    in_code = False
    code_lines: list[str] = []
    active_num_id: int | None = None
    reference_mode = False
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        is_numbered = bool(re.match(r"^\d+\. ", line))
        if not is_numbered:
            active_num_id = None
        if line == "[[FIGURE_DR_VOM]]":
            add_figure(doc)
            i += 1
            continue
        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, parse_table(block))
            continue
        if line.startswith("> "):
            add_callout(doc, line[2:])
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:], size=12, color=NAVY)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:], size=16, color=BLUE)
            reference_mode = line[3:].strip() == "参考来源"
        elif line.startswith("# "):
            pass
        elif is_numbered:
            if active_num_id is None:
                active_num_id = new_numbering_instance(doc)
            p = add_body_paragraph(
                doc,
                re.sub(r"^\d+\. ", "", line),
                style="List Number",
                size=8.9 if reference_mode else 11,
            )
            set_numbering(p, active_num_id)
            if reference_mode:
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.0
        elif line.startswith("- "):
            add_body_paragraph(doc, line[2:], style="List Bullet")
        else:
            p = add_body_paragraph(doc, line, size=9 if reference_mode else 11)
            if reference_mode:
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
        i += 1


def audit_geometry(doc: Document) -> None:
    section = doc.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.left_margin == Inches(1)
    assert section.right_margin == Inches(1)
    for table in doc.tables:
        widths = []
        for col in table._tbl.tblGrid.gridCol_lst:
            widths.append(int(col.get(qn("w:w"))))
        assert sum(widths) == 9360, widths
        for row in table.rows:
            assert len(row.cells) == len(widths)


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    make_chart()
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_markdown_body(doc)
    audit_geometry(doc)
    props = doc.core_properties
    props.title = "ZDS 论文拒稿复盘与 DR-VOM 最终实验报告"
    props.subject = "MITRE ATT&CK next-technique prediction evidence audit"
    props.author = "Research audit"
    props.keywords = "MITRE ATT&CK, DR-VOM, domain shift, Markov, reproducibility"
    props.comments = "Generated from the audited Markdown report; evidence cutoff 2026-07-30."
    doc.save(OUTPUT)
    print(OUTPUT)
    print(CHART)


if __name__ == "__main__":
    build()
